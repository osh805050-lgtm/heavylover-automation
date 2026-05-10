#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
setup_proposal_folder.py — 사업계획서 final 단계 자동 폴더·서류 셋업

호출: python tools/setup_proposal_folder.py {사업명} {YYYY-MM-DD}
예  : python tools/setup_proposal_folder.py 강한소상공인 2026-05-03

기능:
  1. {TARGET_BASE}/{날짜} {사업명}/ 폴더 생성 (제출서류·심사자료·산출물_md 하위)
     - TARGET_BASE 기본값: ~/OneDrive/헤비로버_제출 (repo 외부 — 개인 신분 서류 git 격리)
     - 환경변수 HEAVYLOVER_SUBMIT_BASE 로 override 가능
  2. proposals/outputs/{사업명}-{날짜}-*.md 12개 → 산출물_md/ 복사
  3. final.md → {사업명}_사업계획서_{날짜}.docx 변환
     - pandoc 우선 / python-docx fallback (단, 표·체크박스 있으면 fallback 거부)
  4. 제출서류 자동 탐색: proposals/knowledge/submission-manifest.json 기반
     - 매칭 후보 → 사용자 confirm (interactive) / 배치 모드는 unverified/ 격리
     - manifest 미존재 시 graceful fallback (legacy SUBMISSION_DOCS dict)
  5. 제출_체크리스트.md 생성 (구비 ✅ / 직접준비 ⚠️ / unverified 🔍)

의존성: pandoc (권장) / python-docx (fallback, 표·체크박스 없을 때만)
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path

# ───────────────────────── 경로 정의 ─────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUTS_DIR = PROJECT_ROOT / "proposals" / "outputs"
LESSONS_DIR = PROJECT_ROOT / "proposals" / "lessons"
RUBRIC_JSON = PROJECT_ROOT / "proposals" / "knowledge" / "psst-rubric.json"
MANIFEST_JSON = PROJECT_ROOT / "proposals" / "knowledge" / "submission-manifest.json"


def resolve_target_base() -> Path:
    """제출 폴더 루트 결정. 환경변수 우선 → 기본값(~/OneDrive/헤비로버_제출).

    repo 외부에 두는 이유: 주민등록등본·통장사본 등 개인 신분 서류가
    git history에 남는 사고 방지 (Codex review 2026-05-10 critical).
    """
    override = os.environ.get("HEAVYLOVER_SUBMIT_BASE")
    if override:
        return Path(override).expanduser().resolve()
    return (Path.home() / "OneDrive" / "헤비로버_제출").resolve()


# 제출서류 자동 탐색 위치 (헤비로버 환경)
SEARCH_ROOTS = [
    Path.home() / "OneDrive" / "바탕 화면",
    Path.home() / "OneDrive" / "Documents",
    Path.home() / "OneDrive" / "문서",
    Path.home() / "Documents",
    Path.home() / "Desktop",
]

# Legacy fallback (manifest 미존재 시만 사용)
SUBMISSION_DOCS_LEGACY = {
    "사업자등록증": ["사업자등록증", "business_registration", "사업자_등록증"],
    "통장사본": ["통장사본", "통장_사본", "bankbook", "계좌사본"],
    "4대보험가입자명부": ["4대보험", "가입자명부", "national_pension", "보험가입자"],
    "HACCP인증서": ["HACCP", "haccp", "안전관리인증"],
    "OEM협약서": ["OEM", "oem", "공급협약", "제조위탁"],
    "법인등기부등본": ["등기부등본", "법인등기", "corporate_registry"],
    "임대차계약서": ["임대차", "임대계약", "lease"],
    "주민등록등본": ["주민등록등본", "주민_등본", "resident_registration"],
    "납세증명서": ["납세증명", "납세_증명", "tax_clearance"],
    "신용등급확인서": ["신용등급", "credit_rating"],
}
DEFAULT_EXTENSIONS = ("pdf", "jpg", "jpeg", "png", "hwp", "docx", "xlsx")

# ───────────────────────── 유틸 ─────────────────────────
def log(msg: str, level: str = "INFO") -> None:
    icons = {"INFO": "▶", "OK": "✅", "WARN": "⚠️", "ERR": "❌", "ASK": "❓"}
    print(f"{icons.get(level, '·')} {msg}")


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_manifest() -> dict | None:
    """submission-manifest.json 로드. 미존재·파싱 실패 시 None 반환."""
    if not MANIFEST_JSON.exists():
        log(f"manifest 미존재 → legacy 모드: {MANIFEST_JSON.name}", "WARN")
        return None
    try:
        return json.loads(MANIFEST_JSON.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        log(f"manifest 파싱 실패 → legacy 모드: {e}", "WARN")
        return None


def docs_from_manifest(manifest: dict | None, business: str) -> list[dict]:
    """manifest에서 사업명별 필요 서류 목록 합성. 없으면 legacy dict 변환."""
    if not manifest:
        return [
            {
                "key": k,
                "keywords": v,
                "extensions": list(DEFAULT_EXTENSIONS),
                "required_metadata": [],
                "sensitive": True,
            }
            for k, v in SUBMISSION_DOCS_LEGACY.items()
        ]
    docs = list(manifest.get("default_docs", []))
    program = manifest.get("programs", {}).get(business)
    if program:
        docs.extend(program.get("additional_docs", []))
    # extensions 기본값 채움
    for d in docs:
        if not d.get("extensions"):
            d["extensions"] = list(DEFAULT_EXTENSIONS)
    return docs


# ───────────────────────── 1. 폴더 생성 ─────────────────────────
def create_folder_structure(target_dir: Path) -> dict:
    """제출서류·심사자료·산출물_md·unverified 하위 폴더 생성."""
    subs = {
        "submit": target_dir / "제출서류",
        "review": target_dir / "심사자료",
        "outputs_md": target_dir / "산출물_md",
        "unverified": target_dir / "제출서류" / "unverified",
    }
    for p in subs.values():
        ensure_dir(p)
    log(f"폴더 생성: {target_dir}", "OK")
    return subs


# ───────────────────────── 2. 산출물 복사 ─────────────────────────
def copy_outputs(business: str, date: str, dest: Path) -> list[Path]:
    """proposals/outputs/{business}-{date}-*.md 모두 복사."""
    pattern = f"{business}-{date}-*.md"
    files = sorted(OUTPUTS_DIR.glob(pattern))
    if not files:
        log(f"산출물 없음: {pattern}", "WARN")
        return []
    for f in files:
        shutil.copy2(f, dest / f.name)
    log(f"산출물 {len(files)}개 복사 → 산출물_md/", "OK")
    return files


# ───────────────────────── 3. Word 변환 ─────────────────────────
TABLE_PATTERN = re.compile(r"^\s*\|.*\|\s*$", re.MULTILINE)
TABLE_DIVIDER = re.compile(r"^\s*\|?\s*:?-{2,}", re.MULTILINE)
CHECKBOX_PATTERN = re.compile(r"^\s*[-*]\s+\[[ xX]\]", re.MULTILINE)


def md_has_complex_features(md_text: str) -> tuple[bool, list[str]]:
    """python-docx fallback이 손실시킬 요소(표·체크박스) 감지."""
    found: list[str] = []
    if TABLE_DIVIDER.search(md_text) and TABLE_PATTERN.search(md_text):
        found.append("표 (| --- | divider 감지)")
    if CHECKBOX_PATTERN.search(md_text):
        found.append("체크박스 (- [ ] 감지)")
    return (len(found) > 0, found)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> bool:
    """final.md → docx. pandoc 우선, fallback은 python-docx (표·체크박스 없을 때만).

    중요(Codex review 2026-05-10 high): fallback이 표·서식 손실해도 success 반환
    하던 버그 수정. 표·체크박스 발견 시 명확히 실패 반환.
    """
    if not md_path.exists():
        log(f"변환 대상 없음: {md_path.name}", "ERR")
        return False

    # 1차: pandoc
    try:
        subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path),
             "--from", "gfm", "--to", "docx"],
            check=True, capture_output=True, text=True
        )
        log(f"pandoc 변환 성공 → {docx_path.name}", "OK")
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        log("pandoc 미설치 또는 실패 → fallback 검사 진입", "WARN")

    # fallback 진입 전: 표·체크박스 발견 시 차단
    md_text = md_path.read_text(encoding="utf-8")
    has_complex, features = md_has_complex_features(md_text)
    if has_complex:
        log(f"docx fallback 거부: {', '.join(features)} → 손실 위험.", "ERR")
        log("pandoc 필수: `choco install pandoc` (관리자) 후 재실행", "ERR")
        return False

    # 2차: python-docx fallback (단순 텍스트만 있을 때)
    try:
        from docx import Document
    except ImportError:
        log("python-docx 미설치. `pip install python-docx` 후 재실행", "ERR")
        return False

    save_path = docx_path
    if docx_path.exists():
        from datetime import datetime as _dt
        suffix = _dt.now().strftime("%H%M%S")
        save_path = docx_path.with_stem(docx_path.stem + f"_{suffix}")
        log(f"기존 docx 잠금 감지 → {save_path.name}으로 저장", "WARN")

    try:
        doc = Document()
        for line in md_text.splitlines():
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:].strip(), level=4)
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line)
        doc.save(str(save_path))
        log(f"python-docx 변환 성공 → {save_path.name} (표·체크박스 없음 확인)", "OK")
        return True
    except PermissionError:
        log("docx 저장 권한 오류 — 파일이 열려있으면 닫고 재실행", "ERR")
        return False


# ───────────────────────── 4. 제출서류 자동 탐색 ─────────────────────────
def search_candidates(keywords: list[str], extensions: list[str], limit: int = 5) -> list[Path]:
    """SEARCH_ROOTS 전체에서 키워드+확장자 매칭 후보 수집 (mtime desc)."""
    hits: list[Path] = []
    seen: set[str] = set()
    for root in SEARCH_ROOTS:
        if not root.exists():
            continue
        for kw in keywords:
            for ext in extensions:
                try:
                    matches = list(root.rglob(f"*{kw}*.{ext}"))
                except OSError:
                    # PermissionError·FileNotFoundError 모두 OSError 하위
                    continue
                for m in matches:
                    key = str(m).lower()
                    if key not in seen and m.is_file():
                        seen.add(key)
                        hits.append(m)
    hits.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return hits[:limit]


def confirm_match(doc_key: str, candidate: Path, required_metadata: list[str], interactive: bool) -> bool:
    """사용자 confirm. 비-interactive면 항상 False (unverified 격리)."""
    mtime = datetime.fromtimestamp(candidate.stat().st_mtime).strftime("%Y-%m-%d")
    size_kb = candidate.stat().st_size / 1024
    print()
    log(f"[{doc_key}] 후보: {candidate}", "ASK")
    log(f"        수정일: {mtime}  ·  크기: {size_kb:.1f}KB", "INFO")
    if required_metadata:
        log(f"        검증 항목: {', '.join(required_metadata)}", "INFO")
    if not interactive:
        log("        (비-interactive 모드 → unverified/로 격리)", "WARN")
        return False
    try:
        ans = input("        이 파일이 맞습니까? [y/N]: ").strip().lower()
    except EOFError:
        return False
    return ans in ("y", "yes")


def find_submission_docs(
    docs: list[dict],
    submit_dir: Path,
    unverified_dir: Path,
    interactive: bool,
) -> dict[str, dict]:
    """manifest 정의된 서류별 후보 검색 → 사용자 confirm → 배치/격리.

    반환: {doc_key: {"status": "confirmed|unverified|missing", "path": Path|None, "candidates": [...]}}.
    """
    result: dict[str, dict] = {}
    for d in docs:
        key = d["key"]
        keywords = d.get("keywords", [])
        extensions = d.get("extensions", list(DEFAULT_EXTENSIONS))
        required_metadata = d.get("required_metadata", [])
        candidates = search_candidates(keywords, extensions)

        if not candidates:
            log(f"{key} 후보 없음 — 직접 준비 필요", "WARN")
            result[key] = {"status": "missing", "path": None, "candidates": []}
            continue

        confirmed_path: Path | None = None
        for cand in candidates:
            if confirm_match(key, cand, required_metadata, interactive):
                try:
                    dest = submit_dir / cand.name
                    shutil.copy2(cand, dest)
                    confirmed_path = dest
                    log(f"{key} 확인 완료 → 제출서류/{cand.name}", "OK")
                except Exception as e:
                    log(f"{key} 복사 실패: {e}", "ERR")
                break

        if confirmed_path:
            result[key] = {"status": "confirmed", "path": confirmed_path, "candidates": candidates}
        else:
            # 사용자가 모두 거부 OR 비-interactive 모드 → unverified/에 모든 후보 격리
            unverified_paths: list[Path] = []
            for cand in candidates:
                try:
                    dest = unverified_dir / f"{key}__{cand.name}"
                    shutil.copy2(cand, dest)
                    unverified_paths.append(dest)
                except Exception as e:
                    log(f"{key} unverified 복사 실패 ({cand.name}): {e}", "WARN")
            log(f"{key} unverified 격리: {len(unverified_paths)}건 → 사용자가 직접 검토 후 제출서류/로 이동", "WARN")
            result[key] = {"status": "unverified", "path": None, "candidates": unverified_paths}
    return result


# ───────────────────────── 5. 체크리스트 생성 ─────────────────────────
def generate_checklist(
    target_dir: Path,
    business: str,
    date: str,
    found_docs: dict[str, dict],
    output_files: list[Path],
    docx_ok: bool,
    target_base: Path,
) -> None:
    """제출_체크리스트.md 생성."""
    today = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [
        f"# {business} 제출 체크리스트",
        "",
        f"**생성**: {today}  ·  **사업명**: {business}  ·  **빌드업 날짜**: {date}",
        f"**제출 폴더 루트**: `{target_base}` (repo 외부)",
        "",
        "---",
        "",
        "## 자동 구비된 항목",
        "",
        f"- {'✅' if docx_ok else '❌'} Word 사업계획서 변환본 (`{business}_사업계획서_{date}.docx`)",
        f"- ✅ 산출물 12종 복사 (산출물_md/, 총 {len(output_files)}개)",
        "",
        "## 제출서류 (manifest 기반 자동 탐색 결과)",
        "",
    ]
    for doc_key, info in found_docs.items():
        status = info["status"]
        if status == "confirmed":
            p = info["path"]
            name = p.name if p else "?"
            lines.append(f"- ✅ **{doc_key}** — 사용자 확인 완료: `{name}`")
        elif status == "unverified":
            n = len(info["candidates"])
            lines.append(
                f"- 🔍 **{doc_key}** — 후보 {n}건 unverified/에 격리. "
                "직접 검토 후 맞는 파일을 `제출서류/`로 이동, 나머지 삭제"
            )
        else:
            lines.append(f"- ⚠️ **{doc_key}** — 후보 미발견. 직접 준비 후 `제출서류/`에 업로드")
    lines += [
        "",
        "## 사용자 직접 액션 필수",
        "",
        "1. 🔍 `제출서류/unverified/` 폴더 내 모든 파일 1개씩 열어 확인 → 맞는 것만 상위로 이동",
        "2. ⚠️ 미발견 서류 직접 준비 후 `제출서류/` 폴더에 업로드",
        f"3. ⚠️ `fact-check.md`의 외부 인용 URL 본인이 1회 클릭 검증 (위치: 산출물_md/{business}-{date}-fact-check.md)",
        f"4. ⚠️ `devil-attack.md`의 발표 Q&A 10개 답변 연습 (위치: 산출물_md/{business}-{date}-devil-attack.md)",
        "5. ⚠️ docx 파일 직접 열어 표·서식 깨짐 여부 최종 확인 (pandoc 미사용 시 표·체크박스 자동 차단됨)",
        "6. ⚠️ 신청 후 결과 통지 받으면 `proposals/knowledge/submissions-log/`의 stub에 4칸 (결과·점수·코멘트·회고) 채우기",
        "",
        "## 주의",
        "",
        "- **개인 신분 서류는 git 추적 외부 폴더(`헤비로버_제출/`)에 저장됨** — 절대 repo 내부로 이동 금지.",
        "- 자동 탐색은 키워드+확장자 매칭 → 다른 사업의 옛 서류 매칭 가능성 있음. unverified/ 격리분 반드시 1회 열어 확인.",
        "- pandoc 미설치 + 표/체크박스 포함 시 docx 변환 실패 (의도된 동작). `choco install pandoc` 후 재실행.",
    ]
    (target_dir / "제출_체크리스트.md").write_text("\n".join(lines), encoding="utf-8")
    log("제출_체크리스트.md 생성 완료", "OK")


# ───────────────────────── 메인 ─────────────────────────
def main() -> int:
    parser = argparse.ArgumentParser(description="사업계획서 final 단계 자동 폴더 셋업")
    parser.add_argument("business", help="사업명 (예: 강한소상공인)")
    parser.add_argument("date", help="빌드업 날짜 YYYY-MM-DD")
    parser.add_argument(
        "--non-interactive",
        action="store_true",
        help="사용자 confirm 생략 (CI/배치). 모든 후보를 unverified/로 격리.",
    )
    args = parser.parse_args()

    business = args.business
    date = args.date
    interactive = not args.non_interactive

    target_base = resolve_target_base()
    ensure_dir(target_base)
    log(f"제출 폴더 루트: {target_base} (repo 외부 — 개인 신분 서류 git 격리)", "INFO")

    target_dir = target_base / f"{date} {business}"
    log(f"대상 폴더: {target_dir}", "INFO")

    # 1. 폴더 구조
    subs = create_folder_structure(target_dir)

    # 2. 산출물 복사
    output_files = copy_outputs(business, date, subs["outputs_md"])

    # 3. Word 변환
    final_md = OUTPUTS_DIR / f"{business}-{date}-final.md"
    docx_path = target_dir / f"{business}_사업계획서_{date}.docx"
    docx_ok = convert_md_to_docx(final_md, docx_path)

    # 4. 제출서류 manifest 기반 탐색
    manifest = load_manifest()
    docs = docs_from_manifest(manifest, business)
    log(f"매칭 대상 서류: {len(docs)}종 (manifest={'O' if manifest else 'X-legacy'})", "INFO")
    found_docs = find_submission_docs(docs, subs["submit"], subs["unverified"], interactive)

    # 5. 체크리스트 생성
    generate_checklist(target_dir, business, date, found_docs, output_files, docx_ok, target_base)

    log("완료", "OK")
    log(f"폴더 열기: explorer \"{target_dir}\"", "INFO")
    return 0


if __name__ == "__main__":
    sys.exit(main())
